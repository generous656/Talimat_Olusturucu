from flask import Flask, render_template, request, redirect, flash, url_for, send_file
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import datetime
import os
import webbrowser

# Açmak istediğiniz URL
url = "http://127.0.0.1:5000/"

# Varsayılan tarayıcıda URL'yi aç
webbrowser.open(url)

app = Flask(__name__)
app.secret_key = 'supersecretkey'

# Örnek hesaplar ve alıcılar listesi
hesaplar = []
alicilar = []
templates_file = 'templates.txt'

def save_templates(templates_file, templates):
    with open(templates_file, 'w') as file:
        for template in templates:
            file.write(template + '\n')

def load_templates(templates_file):
    try:
        with open(templates_file, 'r') as file:
            templates = [line.strip() for line in file.readlines()]
    except FileNotFoundError:
        flash('Şablon dosyası bulunamadı. Lütfen şablon ekleyin.')
        templates = []
    return templates

sablonlar = load_templates(templates_file)

@app.route('/')
def index():
    return render_template('index.html', hesaplar=hesaplar, alicilar=alicilar, sablonlar=sablonlar)

@app.route('/add_account', methods=['GET', 'POST'])
def add_account():
    if request.method == 'POST':
        sirket_adi = request.form['sirket_adi']
        vkn = request.form['vkn']
        banka_adi = request.form['banka_adi']
        banka_subesi = request.form['banka_subesi']
        iban = request.form['iban']
        para_birimi = request.form['para_birimi']
        hesaplar.append((sirket_adi, vkn, banka_adi, banka_subesi, iban, para_birimi))
        save_accounts()
        flash("Yeni hesap başarıyla eklendi!")
        return redirect(url_for('index'))
    return render_template('add_account.html')

@app.route('/add_recipient', methods=['GET', 'POST'])
def add_recipient():
    if request.method == 'POST':
        alici_sirket_adi = request.form['alici_sirket_adi']
        alici_firma_adres = request.form['alici_firma_adres']
        alici_banka = request.form['alici_banka']
        alici_banka_adres = request.form['alici_banka_adres']
        alici_iban = request.form['alici_iban']
        alici_swift = request.form['alici_swift']
        alici_para_birimi = request.form['alici_para_birimi']
        alicilar.append((alici_sirket_adi, alici_firma_adres, alici_banka, alici_banka_adres, alici_iban, alici_swift, alici_para_birimi))
        save_recipients()
        flash("Yeni alıcı başarıyla eklendi!")
        return redirect(url_for('index'))
    return render_template('add_recipient.html')

@app.route('/create_instruction', methods=['POST'])
def create_instruction():
    hesap_index = int(request.form['hesap'])
    alici_index = int(request.form['alici'])
    sablon_index = int(request.form['sablon'])
    tutar = request.form['tutar']
    referans = request.form['referans']
    timestamp = datetime.datetime.now().strftime("%Y%m%d%H%M%S")

    if hesap_index < 0 or hesap_index >= len(hesaplar) or alici_index < 0 or alici_index >= len(alicilar):
        flash('Lütfen geçerli bir hesap ve alıcı seçin.')
        return redirect(url_for('index'))

    if sablon_index < 0 or sablon_index >= len(sablonlar):
        flash('Lütfen geçerli bir şablon seçin.')
        return redirect(url_for('index'))

    hesap = hesaplar[hesap_index]
    alici = alicilar[alici_index]
    tarih = datetime.datetime.now().strftime("%d.%m.%Y")

    doc = Document()
    doc.add_heading('Talimat', 0)
    doc.add_paragraph(f"Tarih: {tarih}\n")
    doc.add_paragraph(sablonlar[sablon_index].format(hesap_sirket_adi=hesap[0], hesap_vkn=hesap[1], hesap_banka_adi=hesap[2], hesap_banka_subesi=hesap[3], hesap_iban=hesap[4], hesap_para_birimi=hesap[5], alici_sirket_adi=alici[0], alici_firma_adres=alici[1], alici_banka=alici[2], alici_banka_adres=alici[3], alici_iban=alici[4], alici_swift=alici[5], alici_para_birimi=alici[6]))
    doc.add_heading('Alıcı Bilgileri:', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Firma Adı             : ").bold = True
    p1.add_run(alici[0]).bold = False

    p2 = doc.add_paragraph()
    p2.add_run("Firma Adresi       : ").bold = True
    p2.add_run(alici[1]).bold = False

    p3 = doc.add_paragraph()
    p3.add_run("Banka Adı             : ").bold = True
    p3.add_run(alici[2]).bold = False

    p4 = doc.add_paragraph()
    p4.add_run("Banka Adresi      : ").bold = True
    p4.add_run(alici[3]).bold = False

    p5 = doc.add_paragraph()
    p5.add_run("IBAN                       : ").bold = True
    p5.add_run(alici[4]).bold = False

    p6 = doc.add_paragraph()
    p6.add_run("SWIFT Kodu        : ").bold = True
    p6.add_run(alici[5]).bold = False

    p7 = doc.add_paragraph()
    p7.add_run("Ödeme Tutarı     : ").bold = True
    p7.add_run(f"{tutar} {hesap[5]}").bold = False

    doc.add_heading('Referans Bilgileri:', level=1)
    doc.add_paragraph(f"{referans}\n")

    doc.add_paragraph(f"Saygılarımızla,\n{hesap[0]}")

    file_path = f"talimatlar\Talimat_{alici[0]}_{tutar}_{hesap[5]}.docx"
    doc.save(file_path)

    flash(f"Talimat başarıyla oluşturuldu: {file_path}")
    return redirect(url_for('download_file', filename=file_path))

@app.route('/download/<filename>')
def download_file(filename):
    return send_file(filename, as_attachment=True)

@app.route('/templates_info')
def templates_info():
    return render_template('templates_info.html', sablonlar=sablonlar)

@app.route('/manage_templates')
def manage_templates():
    return render_template('manage_templates.html', sablonlar=sablonlar)

@app.route('/add_template', methods=['POST'])
def add_template():
    yeni_sablon = request.form['sablon']
    sablonlar.append(yeni_sablon)
    save_templates(templates_file, sablonlar)
    flash("Yeni şablon başarıyla eklendi!")
    return redirect(url_for('manage_templates'))

@app.route('/delete_template', methods=['POST'])
def delete_template():
    index = int(request.form['index'])
    if 0 <= index < len(sablonlar):
        sablonlar.pop(index)
        save_templates(templates_file, sablonlar)
        flash("Şablon başarıyla silindi!")
    else:
        flash("Geçersiz şablon indexi!")
    return redirect(url_for('manage_templates'))

def save_accounts():
    with open('hesaplar.txt', 'w') as file:
        for hesap in hesaplar:
            file.write(','.join(hesap) + '\n')

def load_accounts():
    if os.path.exists('hesaplar.txt'):
        with open('hesaplar.txt', 'r') as file:
            return [line.strip().split(',') for line in file.readlines()]
    else:
        return []

def save_recipients():
    with open('alicilar.txt', 'w') as file:
        for alici in alicilar:
            file.write(','.join(alici) + '\n')

def load_recipients():
    if os.path.exists('alicilar.txt'):
        with open('alicilar.txt', 'r') as file:
            return [line.strip().split(',') for line in file.readlines()]
    else:
        return []

if __name__ == '__main__':
    hesaplar = load_accounts()
    alicilar = load_recipients()
    app.run(debug=True)
